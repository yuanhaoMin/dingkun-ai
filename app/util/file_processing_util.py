import re
from datetime import datetime
from io import BytesIO
from typing import Iterator, List
from langchain.embeddings.openai import OpenAIEmbeddings
import chardet
import docx2txt
from docx import Document as DocxDocument
from fastapi import UploadFile
from langchain.document_loaders import PyPDFLoader, Blob
from langchain.embeddings import OpenAIEmbeddings
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pdfminer.high_level import extract_text
from pymilvus import MilvusClient
from app.config.api_config import get_milvus_uri, get_milvus_token


def detect_encoding(file_path):
    with open(file_path, 'rb') as f:
        result = chardet.detect(f.read())
    return result['encoding']


def process_docx_to_str(file: UploadFile) -> str:
    raw_text = docx2txt.process(file)
    processed_text = remove_blank_lines(raw_text)
    return processed_text


def process_pdf_to_str(file: UploadFile) -> str:
    # Read the content of the uploaded file into a byte stream
    file_content = file.file.read()
    byte_stream = BytesIO(file_content)

    # Extract text from the byte stream
    raw_text = extract_text(byte_stream)
    processed_text = remove_blank_lines(raw_text)

    return processed_text


def remove_blank_lines(input_str: str) -> str:
    # Split the string into lines, filter out blank lines, and join the non-blank lines back into a string
    lines = input_str.splitlines()
    non_blank_lines = [line for line in lines if line.strip()]
    return "\n".join(non_blank_lines)


# enc = tiktoken.get_encoding("cl100k_base")


# def length_function(text: str) -> int:
#     return len(enc.encode(text))


def split_content_into_chunks(text):
    # 参考官网链接：https://python.langchain.com/docs/modules/data_connection/document_transformers/text_splitters/character_text_splitter
    text_spliter = RecursiveCharacterTextSplitter(separators=["\n\n", "\n", " ", ""],
                                                  chunk_size=500,
                                                  chunk_overlap=80,
                                                  # length_function=length_function
                                                  )
    chunks = text_spliter.split_text(text)
    return chunks


def process_pdf_file(file: str):
    # 加载PDF文档
    loader = PyPDFLoader(file.name)
    documents = loader.load()

    # 打印文件名来帮助调试
    print(f"Processing file: {file.name}")

    # 下面是正则匹配，用于找出文件名
    pattern = r"[\\/](?P<filename>[^\\/]+)$"
    match = re.search(pattern, file.name)

    # 检查匹配结果
    if match:
        file_name = match.group("filename")
    else:
        file_name = "Unknown"  # 或者你可以选择其他默认文件名或者抛出异常

    return documents, file_name


def extract_from_docx(docx_path):
    doc = DocxDocument(docx_path)

    # Extract basic metadata
    core_properties = doc.core_properties
    metadata = {
        'title': core_properties.title,
        'author': core_properties.author,
        'created': core_properties.created,
        'modified': core_properties.modified
    }

    # Extract text
    full_text = '\n'.join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    return Document(page_content=full_text, metadata=metadata)


def process_txt_to_str(file: UploadFile) -> str:
    encoding = detect_encoding(file.filename)
    content = file.read().decode(encoding)
    return remove_blank_lines(content)


def create_document_from_file(file: UploadFile,user_id: str) -> List[Document]:
    # Determine file type and process accordingly
    if file.filename.endswith('.txt'):
        content = process_txt_to_str(file)
    elif file.filename.endswith('.pdf'):
        content = process_pdf_to_str(file)
    elif file.filename.endswith('.docx'):
        content = process_docx_to_str(file)
    else:
        raise ValueError("Unsupported file type")

    # Prepare the metadata
    creation_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    metadata = {
        "source": file.filename,
        "creation_time": creation_time,
        "user_id": user_id,
        "page": 0
    }

    # Create a Document instance
    doc = Document(page_content=content, metadata=metadata)

    # Return the Document object inside a list
    return [doc]


def process_and_store_file_to_database(file: UploadFile, user_id: str ,collection_name: str, chunk_size: int = 500, chunk_overlap: int = 100, uri: str = None, token: str = None):
    if uri is None:
        uri = get_milvus_uri()
    if token is None:
        token = get_milvus_token()

    # Step 1: Create a Document from the uploaded file
    doc = create_document_from_file(file,user_id)
    text_spli = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )
    docs = text_spli.split_documents(doc)

    # Extract texts and metadatas from docs
    texts = [document.page_content for document in docs]
    metadatas = [document.metadata for document in docs]

    # Initialize the OpenAIEmbeddings instance
    embedding = OpenAIEmbeddings()

    # Get embeddings for each text
    vectors = [embedding.embed_query(text) for text in texts]

    # Prepare data for insertion
    list_of_rows = []
    for text, metadata, vector in zip(texts, metadatas, vectors):
        row = {
            "source": metadata["source"],
            "creation_time": metadata["creation_time"],
            "page": metadata["page"],
            "text": text,
            "vector": vector
        }
        list_of_rows.append(row)

    # Create MilvusClient instance and insert data
    client = MilvusClient(
        uri=get_milvus_uri(),
        token=get_milvus_token()
    )
    client.insert(collection_name, list_of_rows)
